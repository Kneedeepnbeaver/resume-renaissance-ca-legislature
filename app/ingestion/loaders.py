"""Document loaders for PDF, DOCX, and TXT."""
from io import BytesIO
from pathlib import Path
from typing import Iterator

from pypdf import PdfReader
from docx import Document

from .sanitizer import sanitize_resume_text


def load_uploaded_file(uploaded_file) -> str:
    """Extract text from an uploaded file (Streamlit UploadedFile). Supports PDF, DOCX, TXT, MD."""
    name = getattr(uploaded_file, "name", "file")
    suffix = Path(name).suffix.lower()
    data = uploaded_file.getvalue()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(data))
        parts = [p.extract_text() for p in reader.pages if p.extract_text()]
        return "\n\n".join(parts)
    if suffix in (".docx", ".doc"):
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if suffix in (".txt", ".md", ".markdown"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {suffix}. Use PDF, DOCX, TXT, or MD.")


def _load_pdf(file_path: Path) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _load_docx(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def _load_txt(file_path: Path) -> str:
    """Load plain text or markdown file."""
    return file_path.read_text(encoding="utf-8", errors="replace")


def load_document(file_path: Path, doc_type: str) -> tuple[str, dict]:
    """Load a single document and return (text, metadata)."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        text = _load_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        text = _load_docx(file_path)
    elif suffix in (".txt", ".md", ".markdown"):
        text = _load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    metadata = {
        "source": file_path.name,
        "doc_type": doc_type,
    }
    cleaned = text.strip()
    # Reduce hallucinations: tailored resumes often contain meta-text and target-office lines.
    # Sanitizing keeps only resume content to improve RAG precision.
    if doc_type == "resume":
        cleaned = sanitize_resume_text(cleaned)
        metadata["sanitized"] = True
    return cleaned, metadata


def _infer_doc_type(file_path: Path, sources_dir: Path) -> str:
    """Infer doc_type from folder structure."""
    try:
        rel = file_path.resolve().relative_to(sources_dir.resolve())
        parts = rel.parts
        if len(parts) >= 1:
            folder = parts[0].lower()
            if folder == "resumes":
                return "resume"
            if folder == "experiences":
                return "experience"
            if folder == "books":
                return "book"
    except ValueError:
        pass
    return "experience"  # default for files in sources/ root


def load_documents_from_dir(sources_dir: Path) -> Iterator[tuple[str, dict]]:
    """Load all supported documents from a directory (recursive)."""
    sources_dir = Path(sources_dir)
    if not sources_dir.exists():
        return

    supported = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}
    for file_path in sources_dir.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in supported:
            doc_type = _infer_doc_type(file_path, sources_dir)
            try:
                text, metadata = load_document(file_path, doc_type)
                if text:
                    yield text, metadata
            except Exception as e:
                # Log but continue
                print(f"Warning: Could not load {file_path}: {e}")
